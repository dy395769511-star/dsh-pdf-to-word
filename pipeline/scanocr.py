#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Scan-mode page analysis via PaddleOCR 3.x (layout + text + table recognition).

Result of analyze_page():
  {
    "lines":  [{"text","box":[x1,y1,x2,y2],"score"}],      # recognized text lines
    "regions":[{"label","box":[x1,y1,x2,y2],"score"}],     # layout regions
    "tables": [html_string, ...],                          # table structure html (ordered like table regions)
  }
"""
import os
import pathlib
import re
import sys

sys.stdout.reconfigure(encoding="utf-8")

# Models already cached under ~/.paddlex/official_models (shared cache dir).
MODEL_KWARGS = {
    "text_detection_model_name": "PP-OCRv5_mobile_det",
    "text_recognition_model_name": "PP-OCRv5_mobile_rec",
    "layout_detection_model_name": "PP-DocLayout_plus-L",
    "table_classification_model_name": "PP-LCNet_x1_0_table_cls",
    "wired_table_structure_recognition_model_name": "SLANeXt_wired",
    "wired_table_cells_detection_model_name": "RT-DETR-L_wired_table_cell_det",
}


class OcrEngine:
    _ocr = None

    @classmethod
    def get(cls):
        if cls._ocr is None:
            # Resolve the PaddleX model cache dir, portably:
            #   $PADDLE_PDX_CACHE_HOME env > <plugin>/.paddlex-cache (bundled)
            #   > paddle default ~/.paddlex (models download on first use).
            cache = os.environ.get("PADDLE_PDX_CACHE_HOME")
            if not cache:
                bundled = pathlib.Path(__file__).resolve().parent.parent / ".paddlex-cache"
                cache = str(bundled) if bundled.is_dir() else os.path.join(
                    os.path.expanduser("~"), ".paddlex")
            os.environ.setdefault("PADDLE_PDX_CACHE_HOME", cache)
            from paddleocr import PPStructureV3
            cls._ocr = PPStructureV3(
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=True,
                use_table_recognition=True,
                # Models for these optional submodules are not in the local
                # cache; layout comes from PP-DocLayout_plus-L (LayoutDetection).
                use_region_detection=False,
                use_seal_recognition=False,
                use_formula_recognition=False,
                # The installed paddle 3.x PIR+oneDNN CPU path crashes with
                # "ConvertPirAttribute2RuntimeAttribute not support
                # [pir::ArrayAttribute<pir::DoubleAttribute>]" on layout
                # detection; plain CPU ("paddle" run mode) works.
                enable_mkldnn=False,
                **MODEL_KWARGS,
            )
        return cls._ocr


def _get(obj, key, default=None):
    if obj is None:
        return default
    if hasattr(obj, "get"):
        try:
            v = obj.get(key)
            if v is not None:
                return v
        except Exception:
            pass
    try:
        v = obj[key]
        if v is not None:
            return v
    except Exception:
        pass
    try:
        v = getattr(obj, key, None)
        if v is not None:
            return v
    except Exception:
        pass
    return default


def _poly_to_box(b):
    """Normalize [x1,y1,x2,y2] / flat 8+ coords / (N,2) point list to a box."""
    try:
        b = list(b)
    except Exception:
        return None
    if len(b) == 4:
        try:
            return [float(b[0]), float(b[1]), float(b[2]), float(b[3])]
        except (TypeError, ValueError):
            pass
        try:
            pts = [(float(v[0]), float(v[1])) for v in b]
        except (TypeError, ValueError, IndexError):
            return None
        xs = [p[0] for p in pts]
        ys = [p[1] for p in pts]
        return [min(xs), min(ys), max(xs), max(ys)]
    if len(b) >= 8:
        try:
            xs = [float(b[j]) for j in range(0, len(b) - 1, 2)]
            ys = [float(b[j]) for j in range(1, len(b), 2)]
        except (TypeError, ValueError):
            return None
        return [min(xs), min(ys), max(xs), max(ys)]
    return None


def _first_present(*cands):
    """First non-None, non-empty candidate. Safe for numpy arrays, whose
    truth value is ambiguous, so use identity/length checks only."""
    for c in cands:
        if c is None:
            continue
        try:
            if len(c) > 0:
                return c
        except TypeError:
            return c
    return []


def analyze_page(img_path, ocr_engine):
    res = ocr_engine.predict(img_path)
    page = None
    for r in res:
        page = r
        break
    if page is None:
        return {"lines": [], "regions": [], "tables": [], "keys": []}

    keys = []
    try:
        keys = list(page.keys())
    except Exception:
        pass

    ocr_res = _get(page, "overall_ocr_res") or {}
    texts = _first_present(_get(ocr_res, "rec_texts"))
    boxes = _first_present(_get(ocr_res, "rec_boxes"), _get(ocr_res, "rec_polys"))
    scores = _first_present(_get(ocr_res, "rec_scores"))
    lines = []
    for i, t in enumerate(texts):
        t = (t or "").strip()
        if not t:
            continue
        b = boxes[i] if i < len(boxes) else None
        if b is None:
            continue
        box = _poly_to_box(b)
        if box is None:
            continue
        sc = scores[i] if i < len(scores) else 1.0
        lines.append({"text": t, "box": box, "score": float(sc)})

    layout = _get(page, "layout_det_res")
    regions = []
    if isinstance(layout, dict):
        raw_boxes = layout.get("boxes") or []
    else:
        raw_boxes = []
    for rb in raw_boxes:
        if isinstance(rb, dict):
            box = rb.get("coordinate") or rb.get("box")
            label = rb.get("label") or rb.get("name") or "text"
            score = rb.get("score", 1.0)
        else:
            box = rb
            label = "text"
            score = 1.0
        if not box:
            continue
        box = [float(box[0]), float(box[1]), float(box[2]), float(box[3])]
        regions.append({"label": str(label), "box": box, "score": float(score)})

    tables = _get(page, "table_res_list") or []
    norm = []
    for t in tables:
        if not t:
            continue
        if isinstance(t, str):
            norm.append({"html": t, "cell_boxes": None})
            continue
        h = t.get("pred_html") if isinstance(t, dict) else None
        if isinstance(h, dict):
            h = h.get("html")
        if not (isinstance(h, str) and h):
            continue
        cell_boxes = []
        for cb in t.get("cell_box_list") or []:
            b = _poly_to_box(cb)
            if b:
                cell_boxes.append(b)
        norm.append({"html": h, "cell_boxes": cell_boxes or None})
    tables = norm

    images = []
    for im in _get(page, "imgs_in_doc") or []:
        if not isinstance(im, dict):
            continue
        coord = im.get("coordinate")
        img_obj = im.get("img")
        if not coord or img_obj is None:
            continue
        box = _poly_to_box(coord)
        if box is None:
            continue
        images.append({
            "box": box,
            "img": img_obj,
            "label": str(im.get("label") or "image"),
            "score": float(im.get("score") or 0.0),
        })

    return {
        "lines": lines,
        "regions": regions,
        "tables": tables,
        "images": images,
        "keys": keys,
    }


# ---------------------------------------------------------------------------
# docx building from OCR result
# ---------------------------------------------------------------------------
STANDARD_SIZES = [9, 10.5, 12, 14, 16, 18, 22, 26, 28]


def nearest_std(v):
    return min(STANDARD_SIZES, key=lambda s: abs(s - v))


def _rects_overlap(a, b, pad=3.0):
    return not (a[2] + pad < b[0] or b[2] + pad < a[0] or a[3] + pad < b[1] or b[3] + pad < a[1])


def _center_in(pt, r, pad=3.0):
    return (r[0] - pad) <= pt[0] <= (r[2] + pad) and (r[1] - pad) <= pt[1] <= (r[3] + pad)


def _int_attr(attrs, name, default=1):
    m = re.search(name + r'\s*=\s*"?(\d+)', attrs or "", re.I)
    if not m:
        return default
    try:
        return int(m.group(1))
    except ValueError:
        return default


def parse_html_table(html):
    """Parse recognized table HTML into rows of cell dicts
    {"text", "col" (colspan), "row" (rowspan), "h" (is <th>)}."""
    rows = re.findall(r"<tr[^>]*>(.*?)</tr>", html or "", re.S | re.I)
    out = []
    for row in rows:
        cells = []
        for m in re.finditer(r"<(t[dh])\b([^>]*)>(.*?)</t[dh]>", row, re.S | re.I):
            tag, attrs, inner = m.group(1), m.group(2) or "", m.group(3)
            text = re.sub(r"<[^>]+>", "", inner).replace("&nbsp;", " ")
            text = re.sub(r"\s+", " ", text).strip()
            cells.append({
                "text": text,
                "col": max(1, _int_attr(attrs, "colspan")),
                "row": max(1, _int_attr(attrs, "rowspan")),
                "h": tag.lower() == "h",
            })
        if cells:
            out.append(cells)
    return out


def group_paragraphs(lines, page_w, page_h):
    """Cluster OCR lines into paragraphs (with simple column detection)."""
    if not lines:
        return []
    import collections
    xs = collections.Counter()
    for l in lines:
        xs[round(l["box"][0], -1)] += 1
    # detect column split: a wide empty vertical band
    if lines:
        min_x0 = min(l["box"][0] for l in lines)
        max_x1 = max(l["box"][2] for l in lines)
        band_w = (max_x1 - min_x0)
        mid = None
        if band_w > page_w * 0.5:
            # find a gap with no line boxes
            covered = sorted(set(int(x) for l in lines for x in (l["box"][0], l["box"][2])))
            prev = covered[0] if covered else 0
            best_gap = 0
            for x in covered[1:]:
                gap = x - prev
                if gap > best_gap and min_x0 + band_w * 0.25 < prev + gap / 2 < max_x1 - band_w * 0.25:
                    best_gap = gap
                    mid = prev + gap / 2
                prev = x
            if mid and best_gap > page_w * 0.04:
                left = [l for l in lines if (l["box"][0] + l["box"][2]) / 2 < mid]
                right = [l for l in lines if (l["box"][0] + l["box"][2]) / 2 >= mid]
                return _cluster(left) + _cluster(right)
    return _cluster(lines)


def _cluster(lines):
    lines = sorted(lines, key=lambda l: (l["box"][1], l["box"][0]))
    paras = []
    for l in lines:
        placed = False
        for para in paras[-3:]:
            last = para["lines"][-1]
            lh = max(6.0, last["box"][3] - last["box"][1])
            gap = l["box"][1] - last["box"][3]
            dx = abs(l["box"][0] - last["box"][0])
            if gap < lh * 1.4 and dx < max(24.0, lh * 2.5):
                para["lines"].append(l)
                placed = True
                break
        if not placed:
            paras.append({"lines": [l]})
    return paras


def _line_height(l):
    return max(6.0, l["box"][3] - l["box"][1])


def add_scan_page(doc, data, page_img_path, page_rect, body_size, page_idx, warnings,
                  body_is_global=False):
    """Build docx content for one scanned page. Returns stats."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = data["lines"]
    regions = data["regions"]
    tables_html = data["tables"]
    images = data.get("images") or []
    if data.get("keys") and warnings is not None and page_idx == 0:
        warnings.append("paddleocr result keys: %s" % ",".join(str(k) for k in data["keys"]))

    # OCR boxes are in rendered-image pixels; convert everything to pt so
    # layout math (sizes, gaps, table widths) is in document units.
    from PIL import Image
    try:
        W_px = Image.open(page_img_path).size[0]
        s = page_rect.width / float(W_px) if W_px else 1.0
    except Exception:
        s = 1.0
    if s != 1.0:
        for coll in (lines, regions, images):
            for it in coll:
                it["box"] = [v * s for v in it["box"]]
        for t in tables_html:
            if isinstance(t, dict) and t.get("cell_boxes"):
                t["cell_boxes"] = [[v * s for v in b] for b in t["cell_boxes"]]

    region_by_kind = {"table": [], "figure": [], "image": [], "chart": [],
                      "title": [], "text": [], "skip": [], "other": []}
    for r in regions:
        lab = r["label"].lower()
        if "table" in lab:
            region_by_kind["table"].append(r)
        elif "figure" in lab or "image" in lab or "chart" in lab or "photo" in lab:
            region_by_kind["figure"].append(r)
        elif "title" in lab:
            region_by_kind["title"].append(r)
        elif "text" in lab or "paragraph" in lab or lab == "":
            region_by_kind["text"].append(r)
        elif any(k in lab for k in ("header", "footer", "page_number", "page number", "footnote", "seal", "formula", "equation")):
            region_by_kind["skip"].append(r)
        else:
            region_by_kind["other"].append(r)

    table_regions = region_by_kind["table"]
    figure_regions = region_by_kind["figure"]
    title_regions = region_by_kind["title"]

    def in_region(box, regions_, pad=6.0):
        c = ((box[0] + box[2]) / 2, (box[1] + box[3]) / 2)
        for r in regions_:
            if _center_in(c, r["box"], pad=pad):
                return True
        return False

    free_lines = [l for l in lines
                  if not in_region(l["box"], table_regions)
                  and not in_region(l["box"], figure_regions)]
    title_lines = [l for l in free_lines if in_region(l["box"], title_regions)]
    body_lines = [l for l in free_lines if not in_region(l["box"], title_regions)]

    # Estimate the body size from this page's free OCR line heights (pt).
    # Table/figure lines are excluded: cell boxes are cell sizes, not text
    # heights. OCR boxes run ~1.2-1.4x the font size, so apply a 0.85
    # correction before snapping to a standard size. When a document-level
    # running mode is already established (body_is_global) it is honored;
    # otherwise the page's own mode is used. The passed-in body_size is the
    # fallback. Raw votes are returned so the caller can maintain the
    # document-level mode across pages.
    import collections
    _size_votes = collections.Counter()
    for l in free_lines:
        _size_votes[nearest_std(_line_height(l) * 0.85)] += 1
    if not body_is_global and _size_votes:
        body_size = _size_votes.most_common(1)[0][0]
    stats = {"body_size": body_size, "body_votes": dict(_size_votes),
             "paragraphs": 0, "tables": 0, "images": 0}

    paras = group_paragraphs(body_lines, page_rect.width, page_rect.height)
    if not paras:
        paras = group_paragraphs(title_lines + body_lines, page_rect.width, page_rect.height)
        title_lines = []

    # y order of all elements
    elements = []
    for p in paras:
        y0 = min(l["box"][1] for l in p["lines"])
        elements.append({"kind": "para", "y": y0, "para": p})
    for r in table_regions:
        elements.append({"kind": "table", "y": r["box"][1], "region": r})
    for r in figure_regions:
        elements.append({"kind": "figure", "y": r["box"][1], "region": r})
    for r in title_regions:
        tlines = [l for l in title_lines if in_region(l["box"], [r])]
        if not tlines:
            continue
        tlines = sorted(tlines, key=lambda l: (l["box"][1], l["box"][0]))
        elements.append({"kind": "title", "y": tlines[0]["box"][1], "region": r, "para": {"lines": tlines}})
    elements.sort(key=lambda e: e["y"])

    table_i = 0
    for el in elements:
        if el["kind"] in ("para", "title"):
            is_title = el["kind"] == "title"
            texts = [l["text"] for l in el["para"]["lines"]]
            first_t = texts[0] if texts else ""
            cjk = sum(1 for ch in first_t if "\u4e00" <= ch <= "\u9fff")
            joiner = "" if (first_t and cjk * 2 >= len(first_t)) else " "
            text = joiner.join(texts)
            level = 0
            sizes = [_line_height(l) for l in el["para"]["lines"]]
            avg_h = sum(sizes) / max(1, len(sizes))
            # OCR line boxes run ~1.2-1.4x the font size; the 0.85 factor
            # recovers the true size (22/16/14 for title/H1/H2). Absolute
            # tiers mirror the digital mapping (22->H1, 16->H2, 14->H3).
            std = nearest_std(avg_h * 0.85)
            if std >= 20:
                level = 1
            elif std >= 15:
                level = 2
            elif std >= 13 and is_title:
                level = 3
            p = doc.add_paragraph()
            if level:
                try:
                    p.style = doc.styles["Heading %d" % level]
                except Exception:
                    pass
            pf = p.paragraph_format
            cxs = [ (l["box"][0] + l["box"][2]) / 2 for l in el["para"]["lines"]]
            page_cx = page_rect.width / 2
            if len(set(round(c / 8) for c in cxs)) <= 2 and abs(sum(cxs) / len(cxs) - page_cx) < page_rect.width * 0.08:
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            if level:
                run.font.size = Pt(std)
                run.font.bold = True
                from convert import set_run_fonts, map_font
                set_run_fonts(run, "Arial", "黑体")
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                run.font.size = Pt(body_size)
                from convert import set_run_fonts, map_font
                set_run_fonts(run, "Times New Roman", "宋体")
            stats["paragraphs"] += 1
        elif el["kind"] == "table":
            r = el["region"]
            ok = False
            item = tables_html[table_i] if table_i < len(tables_html) else None
            table_i += 1
            html = cell_boxes = None
            if isinstance(item, dict):
                html = item.get("html")
                cell_boxes = item.get("cell_boxes")
            elif isinstance(item, str):
                html = item
            if html:
                grid = parse_html_table(html)
                if grid:
                    ok = _add_scan_table(doc, grid, r, page_rect, body_size,
                                         cell_boxes)
            if not ok:
                # fall back: crop region image
                if _add_region_image(doc, page_img_path, r["box"], page_rect, body_size):
                    ok = True
            if ok:
                stats["tables"] += 1
        elif el["kind"] == "figure":
            if _add_region_image(doc, page_img_path, el["region"]["box"], page_rect, body_size):
                stats["images"] += 1
    return stats


def _add_scan_table(doc, grid, region, page_rect, body_size, cell_boxes=None):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from convert import set_run_fonts

    n_r = len(grid)
    n_c = max(sum(c["col"] for c in row) for row in grid)
    tbl = doc.add_table(rows=n_r, cols=n_c)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tbl._tbl.tblPr.append(layout)
        w_pt = region["box"][2] - region["box"][0]
        # Per-column widths from recognized cell boxes when they give a
        # clean set of n_c+1 boundaries; otherwise uniform.
        col_w = None
        if cell_boxes:
            # Cell boxes jitter a couple of pt; cluster near-duplicate
            # boundaries instead of demanding exact equality.
            xs = sorted({b[0] for b in cell_boxes} | {b[2] for b in cell_boxes})
            tol = max(2.0, w_pt * 0.01)
            groups = []
            for x in xs:
                if groups and x - groups[-1][-1] <= tol:
                    groups[-1].append(x)
                else:
                    groups.append([x])
            bounds = [sum(g) / len(g) for g in groups]
            if len(bounds) == n_c + 1:
                cand = [bounds[i + 1] - bounds[i] for i in range(n_c)]
                if all(w > 0 for w in cand) and abs(sum(cand) - w_pt) < w_pt * 0.3:
                    col_w = cand
        if col_w is None:
            col_w = [w_pt / n_c] * n_c
        for ci, w in enumerate(col_w):
            for ri in range(n_r):
                tbl.cell(ri, ci).width = Pt(w)
        try:
            tblGrid = tbl._tbl.tblGrid
            if tblGrid is not None:
                for gc, w in zip(list(tblGrid), col_w):
                    gc.set(qn("w:w"), str(int(round(w * 20))))
        except Exception:
            pass
    except Exception:
        pass

    # Grid position of each declared cell, honoring colspan/rowspan so a
    # cell under a rowspan lands in the correct column.
    blocked = set()
    positions = []
    for ri, row in enumerate(grid):
        ci = 0
        for c in row:
            while (ri, ci) in blocked:
                ci += 1
            positions.append((ri, ci, c))
            for dr in range(c["row"]):
                for dc in range(c["col"]):
                    if dr or dc:
                        blocked.add((ri + dr, ci + dc))
            ci += c["col"]

    for ri, ci, c in positions:
        if not c["text"]:
            continue
        cellp = tbl.cell(ri, ci)
        para = cellp.paragraphs[0]
        run = para.add_run(c["text"])
        run.font.size = Pt(body_size)
        set_run_fonts(run, "Times New Roman", "宋体")
        if c["h"] or ri == 0:
            run.font.bold = True
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Apply spans only after all text is in place (merge concatenates
    # covered cells, which are empty).
    for ri, ci, c in reversed(positions):
        try:
            if c["col"] > 1:
                tbl.cell(ri, ci).merge(tbl.cell(ri, ci + c["col"] - 1))
            if c["row"] > 1:
                tbl.cell(ri, ci).merge(tbl.cell(ri + c["row"] - 1, ci))
        except Exception:
            pass
    return True


def _add_region_image(doc, page_img_path, region_box, page_rect, body_size):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image

    try:
        img = Image.open(page_img_path)
        W, H = img.size
        scale = W / page_rect.width  # px per pt
        x0 = int(region_box[0] * scale)
        y0 = int(region_box[1] * scale)
        x1 = int(region_box[2] * scale)
        y1 = int(region_box[3] * scale)
        x0 = max(0, min(W, x0)); y0 = max(0, min(H, y0))
        x1 = max(0, min(W, x1)); y1 = max(0, min(H, y1))
        if x1 - x0 < 8 or y1 - y0 < 8:
            return False
        crop = img.crop((x0, y0, x1, y1))
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
        tmp.close()
        crop.convert("RGB").save(tmp.name, quality=88)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        w_pt = region_box[2] - region_box[0]
        run.add_picture(tmp.name, width=Pt(max(30.0, w_pt)))
        import os
        os.unlink(tmp.name)
        return True
    except Exception:
        return False
