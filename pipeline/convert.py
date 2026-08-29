#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""pdf2word pipeline — convert a PDF to DOCX preserving formatting.

Usage:
  python convert.py --pdf <in.pdf> --out <out.docx> [--mode auto|digital|scan]
                    [--assets <dir>] [--dpi 140]

Outputs:
  <out>                  the generated DOCX
  <assets>/pdf_pN.jpg    rendered original PDF pages (for LLM verification)
  <assets>/word_pN.jpg   rendered Word pages (soffice if available, else PIL preview)
  <assets>/convert.json  per-page statistics and warnings

The last stdout line is: PDF2WORD_JSON <json>
"""
import argparse
import io
import json
import os
import re
import shutil
import stat as statmod
import subprocess
import sys
import tempfile
from collections import Counter

import fitz  # PyMuPDF
from PIL import Image

from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")

# ---------------------------------------------------------------------------
# Font mapping: PDF font name -> (ascii font, east-asian font)
# ---------------------------------------------------------------------------
FONT_RULES = [
    ("simsun", "Times New Roman", "宋体"),
    ("nsimsun", "Times New Roman", "宋体"),
    ("stsong", "Times New Roman", "宋体"),
    ("songti", "Times New Roman", "宋体"),
    ("simsun-extb", "Times New Roman", "宋体"),
    ("simhei", "Arial", "黑体"),
    ("stheiti", "Arial", "黑体"),
    ("heiti", "Arial", "黑体"),
    ("simkai", "Arial", "楷体"),
    ("stkaiti", "Arial", "楷体"),
    ("kaiti", "Arial", "楷体"),
    ("fangsong", "Times New Roman", "仿宋"),
    ("stfangsong", "Times New Roman", "仿宋"),
    ("yahei", "Arial", "微软雅黑"),
    ("msyh", "Arial", "微软雅黑"),
    ("dengxian", "Arial", "等线"),
    ("dengxian-light", "Arial", "等线"),
    ("notosanscjk", "Arial", "黑体"),
    ("sourcehansans", "Arial", "黑体"),
    ("notoserifcjk", "Times New Roman", "宋体"),
    ("sourcehanserif", "Times New Roman", "宋体"),
    ("fzstk", "Times New Roman", "宋体"),
    ("fzstzh", "Times New Roman", "宋体"),
    ("fzhst", "Times New Roman", "黑体"),
    ("times", "Times New Roman", "宋体"),
    ("helvetica", "Arial", "宋体"),
    ("arial", "Arial", "宋体"),
    ("courier", "Courier New", "宋体"),
    ("lucida", "Lucida Console", "宋体"),
    ("opensans", "Arial", "微软雅黑"),
    ("microsoftya", "Arial", "微软雅黑"),
    ("calibri", "Calibri", "宋体"),
    ("segoe", "Segoe UI", "宋体"),
]


def map_font(pdf_font_name):
    n = (pdf_font_name or "").lower().replace("-", "").replace("_", "").replace(" ", "")
    for key, ascii_f, ea_f in FONT_RULES:
        if key in n:
            return ascii_f, ea_f
    if pdf_font_name:
        return pdf_font_name, "宋体"
    return "Times New Roman", "宋体"


def set_run_fonts(run, ascii_font, ea_font):
    run.font.name = ascii_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), ea_font)


def color_to_rgb(c):
    return RGBColor((c >> 16) & 0xFF, (c >> 8) & 0xFF, c & 0xFF)


def rect_overlap_v(a, b):
    return max(0.0, min(a[3], b[3]) - max(a[1], b[1]))


def center_in_rect(pt, rect, pad=2.0):
    return (rect[0] - pad) <= pt[0] <= (rect[2] + pad) and (rect[1] - pad) <= pt[1] <= (rect[3] + pad)


# ---------------------------------------------------------------------------
# Digital (text-layer) PDF conversion
# ---------------------------------------------------------------------------
def page_char_count(page):
    d = page.get_text("dict")
    n = 0
    for b in d["blocks"]:
        if b.get("type") == 0:
            for l in b["lines"]:
                for s in l["spans"]:
                    n += len(s.get("text", ""))
    return n, d


def underline_lines(page):
    """Detect underline vector strokes under text lines. Returns set of (y0,y1,x0,x1) line keys."""
    marks = []
    try:
        for dr in page.get_drawings():
            r = dr.get("rect")
            if not r:
                continue
            w = r.width
            h = r.height
            if h < 2.2 and w > 8 and abs(r.height - h) == 0:
                marks.append((round(r.x0, 1), round(r.y0, 1), round(r.x1, 1), round(r.y1, 1)))
    except Exception:
        pass
    return marks


def line_is_underlined(line_bbox, marks):
    for (mx0, my0, mx1, my1) in marks:
        if my0 >= line_bbox[3] - 2.5 and my0 <= line_bbox[3] + 3.0:
            # A real underline sits within the line's own horizontal extent
            # (table rules / separator lines extend well beyond the text line).
            if mx0 >= line_bbox[0] - 2 and mx1 <= line_bbox[2] + 2:
                ov = min(line_bbox[2], mx1) - max(line_bbox[0], mx0)
                if ov >= min(6.0, line_bbox[2] - line_bbox[0]):
                    return True
    return False


def dominant_spans(spans, attr):
    """Return the modal value of attr weighted by char count."""
    c = Counter()
    for s in spans:
        w = max(1, len(s.get("text", "")))
        c[s.get(attr)] += w
    if not c:
        return None
    return c.most_common(1)[0][0]


def fraction(spans, pred):
    total = sum(len(s.get("text", "")) for s in spans) or 1
    hit = sum(len(s.get("text", "")) for s in spans if pred(s))
    return hit / total


def para_info_from_block(block, region, page, ul_marks, page_size):
    """Build paragraph description from a pymupdf text block."""
    lines = block["lines"]
    spans = [s for l in lines for s in l["spans"] if s.get("text", "").strip()]
    total = sum(len(s.get("text", "")) for s in spans) or 1
    if not spans:
        return None

    sizes = [s["size"] for s in spans]
    size = sum(s["size"] * len(s["text"]) for s in spans) / total
    bold = fraction(spans, lambda s: bool(s["flags"] & 16)) > 0.5
    italic = fraction(spans, lambda s: bool(s["flags"] & 2)) > 0.5
    font = dominant_spans(spans, "font")
    color = dominant_spans(spans, "color")
    color = color if isinstance(color, int) else 0

    # line pitch -> line spacing ratio.
    # Word multipliers scale the font's single-line height (about 1.2x the font
    # size for CJK), so divide the measured pitch ratio by 1.2 to keep the
    # docx vertical rhythm close to the PDF (avoids per-page overflow that
    # breaks 1:1 page correspondence).
    ys = [l["bbox"][1] for l in lines]
    pitches = [ys[i + 1] - ys[i] for i in range(len(ys) - 1)]
    pitch = (sum(pitches) / len(pitches)) if pitches else size * 1.2
    ratio = pitch / max(size, 1)
    ls = round(min(2.0, max(0.95, ratio / 1.2)), 2)

    # alignment (center relative to region OR page center)
    region_cx = (region[0] + region[2]) / 2
    page_cx = page_size.width / 2.0
    lcx = (block["bbox"][0] + block["bbox"][2]) / 2
    rl = block["bbox"][2]
    lw = block["bbox"][2] - block["bbox"][0]
    c_tol = max(4.0, 0.02 * (region[2] - region[0]))
    margin_r = page_size.width - region[0]
    r_tol = max(6.0, 0.03 * (region[2] - region[0]))
    if lw < 0.75 * (region[2] - region[0]) and (abs(lcx - region_cx) < c_tol or abs(lcx - page_cx) < c_tol):
        align = WD_ALIGN_PARAGRAPH.CENTER
    elif rl > margin_r - r_tol:
        # right or justified: judge against the symmetric page margin and per-line edges
        reach = [l["bbox"][2] >= margin_r - r_tol for l in lines]
        if len(lines) == 1:
            align = (WD_ALIGN_PARAGRAPH.RIGHT
                     if reach[0] and lw < 0.75 * (region[2] - region[0])
                     else WD_ALIGN_PARAGRAPH.LEFT)
        elif all(reach):
            align = WD_ALIGN_PARAGRAPH.RIGHT
        elif all(reach[:-1]):
            align = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            align = WD_ALIGN_PARAGRAPH.LEFT
    else:
        align = WD_ALIGN_PARAGRAPH.LEFT

    # first line indent
    first_x0 = lines[0]["bbox"][0]
    rest_x0 = min((l["bbox"][0] for l in lines[1:]), default=first_x0)
    first_indent = 0.0
    if first_x0 - rest_x0 > max(4.0, 0.5 * size):
        first_indent = first_x0 - rest_x0

    # underline per line
    ul_lines = set()
    for i, l in enumerate(lines):
        if line_is_underlined(l["bbox"], ul_marks):
            ul_lines.add(i)

    return {
        "bbox": block["bbox"],
        "text": "",
        "lines": lines,
        "size": size,
        "bold": bold,
        "italic": italic,
        "font": font,
        "color": color,
        "align": align,
        "line_spacing": ls,
        "first_indent": first_indent,
        "ul_lines": ul_lines,
        "space_before": 0.0,
        "space_after": 0.0,
    }


def add_digital_paragraph(doc, info, body_size, page):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = info["align"]
    pf.line_spacing = info["line_spacing"]
    if info.get("bg_fill"):
        _apply_paragraph_shading(p, info["bg_fill"])
        if info.get("box_sides") and info.get("bg_stroke"):
            _apply_paragraph_border(p, info["bg_stroke"], info["box_sides"],
                                    info.get("bg_stroke_w") or 0.5)
    if info["first_indent"] > 0:
        pf.first_line_indent = Pt(info["first_indent"])
    if info["space_before"] > 1:
        pf.space_before = Pt(min(36.0, info["space_before"]))
    if info["space_after"] > 1:
        pf.space_after = Pt(min(36.0, info["space_after"]))

    # heading detection
    size = info["size"]
    text = info["text"].strip()
    level = 0
    if text and len(text) <= 60:
        if size >= max(body_size * 1.6, 16):
            level = 1
        elif size >= max(body_size * 1.3, 13.5):
            level = 2
        elif size >= max(body_size * 1.15, 12.5):
            level = 3
    if level:
        try:
            p.style = doc.styles["Heading %d" % level]
        except Exception:
            pass

    # emit runs: merge consecutive spans with identical style
    pending = []

    def flush():
        if not pending:
            return
        run = p.add_run(pending[0]["text"])
        _style_run(run, pending[0], level)
        for extra in pending[1:]:
            run.text += extra["text"]

    for li, line in enumerate(info["lines"]):
        for s in line["spans"]:
            t = s.get("text", "")
            if not t:
                continue
            key = (s["font"], round(s["size"], 1), bool(s["flags"] & 16),
                   bool(s["flags"] & 2), s.get("color", 0), li in info["ul_lines"])
            if pending and key != pending[0]["key"]:
                flush()
                pending = []
            pending.append({"text": t, "key": key, "font": s["font"], "size": s["size"],
                            "bold": bool(s["flags"] & 16), "italic": bool(s["flags"] & 2),
                            "color": s.get("color", 0), "ul": li in info["ul_lines"]})
    flush()
    return p


def _style_run(run, info, level):
    run.font.size = Pt(info["size"])
    # explicit run-level weight from the PDF spans (overrides the Heading
    # style default when the source text is regular weight)
    run.font.bold = info["bold"]
    run.font.italic = info["italic"]
    if info["ul"]:
        run.font.underline = True
    ascii_f, ea_f = map_font(info["font"])
    set_run_fonts(run, ascii_f, ea_f)
    if level:
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif info["color"] not in (0, None):
        run.font.color.rgb = color_to_rgb(info["color"])


def add_digital_table(doc, table, page_spans, body_size):
    """Add one pymupdf Table to the docx. page_spans: list of (bbox, text, style)."""
    try:
        rows = table.rows
        n_r = table.row_count
        n_c = table.col_count
    except Exception:
        return None
    if not rows or not n_r or not n_c:
        return None
    tbl = doc.add_table(rows=n_r, cols=n_c)
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        tblPr = tbl._tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
    except Exception:
        pass

    # column boundaries (x edges) derived from all cell bboxes
    xs = set()
    for row in rows:
        for cell in row.cells:
            if cell is None:
                continue
            cb = cell.bbox if hasattr(cell, "bbox") else cell
            if cb is None:
                continue
            xs.add(round(cb[0], 1))
            xs.add(round(cb[2], 1))
    xs = sorted(xs)
    col_w = [xs[i + 1] - xs[i] for i in range(len(xs) - 1)]
    if len(col_w) != n_c or sum(col_w) <= 0:
        span = (table.bbox[2] - table.bbox[0]) / n_c
        xs = [table.bbox[0] + i * span for i in range(n_c + 1)]
        col_w = [span] * n_c
    for ci, wpt in enumerate(col_w):
        for ri in range(n_r):
            try:
                tbl.cell(ri, ci).width = Pt(wpt)
            except Exception:
                pass
    # keep w:tblGrid in sync with computed column widths (pt -> twips)
    try:
        tblGrid = tbl._tbl.tblGrid
        if tblGrid is not None:
            for ci, gc in enumerate(list(tblGrid)):
                if ci < len(col_w):
                    gc.set(qn("w:w"), str(int(round(col_w[ci] * 20))))
    except Exception:
        pass

    def col_index(x):
        # rightmost boundary index i with xs[i] <= x (column containing x)
        for i in range(len(xs) - 1, 0, -1):
            if x >= xs[i]:
                return i
        return 0

    # place cell text; merge docx columns spanned by a wider PDF cell
    merged = 0
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row.cells):
            if cell is None:
                continue
            cb = cell.bbox if hasattr(cell, "bbox") else cell
            if cb is None:
                continue
            c0 = col_index(cb[0])
            c1 = col_index(cb[2] - 0.01)
            anchor = tbl.cell(ri, c0)
            if c1 > c0:
                try:
                    anchor = anchor.merge(tbl.cell(ri, c1))
                    merged += 1
                except Exception:
                    pass
            # collect lines whose center is inside this cell
            texts = []
            best_style = None
            for (bb, text, style) in page_spans:
                cx = (bb[0] + bb[2]) / 2
                cy = (bb[1] + bb[3]) / 2
                if center_in_rect((cx, cy), cb, pad=1.0) and text.strip():
                    texts.append(text.strip())
                    if best_style is None:
                        best_style = style
            if not texts:
                continue
            first = True
            for t in texts:
                para = anchor.paragraphs[0] if first else anchor.add_paragraph()
                first = False
                run = para.add_run(t)
                run.font.size = Pt(body_size)
                if best_style:
                    run.font.bold = best_style.get("bold", False)
                    ascii_f, ea_f = map_font(best_style.get("font"))
                    set_run_fonts(run, ascii_f, ea_f)
                else:
                    set_run_fonts(run, "Times New Roman", "宋体")
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return {"rows": n_r, "cols": n_c, "merged": merged}


def add_digital_image(doc, pdf_doc, item, region):
    img = None
    xref = item.get("xref")
    if xref:
        try:
            info = pdf_doc.extract_image(xref)
            img = Image.open(io.BytesIO(info["image"]))
            smask = info.get("smask")
            if smask:
                mask = Image.open(io.BytesIO(pdf_doc.extract_image(smask)["image"])).convert("L")
                img = img.convert("RGBA")
                img.putalpha(mask)
        except Exception:
            img = None
    if img is None and item.get("image"):
        try:
            img = Image.open(io.BytesIO(item["image"]))
        except Exception:
            img = None
    if img is None:
        return False
    try:
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tmp.close()
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        img.save(tmp.name)
        w_pt = item["bbox"][2] - item["bbox"][0]
        p = doc.add_paragraph()
        if abs((item["bbox"][0] + item["bbox"][2]) / 2 - (region[0] + region[2]) / 2) < 0.05 * (region[2] - region[0]):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(tmp.name, width=Pt(max(20.0, w_pt)))
        os.unlink(tmp.name)
        return True
    except Exception:
        return False


def _is_real_table(t):
    """Reject pymupdf misdetections where a code block or the whole page text
    area is detected as a table: single-row grids, or a detection containing
    one cell that spans more than half of the table's area (a real table's
    cells form a grid; a misdetected page area is one giant cell)."""
    try:
        if t.row_count < 2 or t.col_count < 2:
            return False
        tw = max(t.bbox[2] - t.bbox[0], 1e-6)
        th = max(t.bbox[3] - t.bbox[1], 1e-6)
        for c in t.cells:
            if not c:
                continue
            ca = max(c[2] - c[0], 0.0) * max(c[3] - c[1], 0.0)
            if ca > 0.5 * tw * th:
                return False
        return True
    except Exception:
        return False


def page_bg_rects(page, min_area=3000.0):
    """Non-white filled rectangles (code-block backgrounds, highlights).

    Returns [(bbox, "RRGGBB" fill, "RRGGBB" stroke or None, stroke_width_pt), ...];
    skips white/near-white fills and rectangles that cover more than 90% of the
    page (page furniture)."""
    out = []
    page_area = page.rect.width * page.rect.height
    try:
        drawings = page.get_drawings()
    except Exception:
        return out
    for dr in drawings:
        fill = dr.get("fill")
        if not fill:
            continue
        r = dr["rect"]
        a = (r[2] - r[0]) * (r[3] - r[1])
        if a < min_area or a > 0.9 * page_area:
            continue
        if all(c > 0.99 for c in fill):
            continue
        fill_hex = "%02X%02X%02X" % (
            int(round(fill[0] * 255)), int(round(fill[1] * 255)),
            int(round(fill[2] * 255)))
        stroke = dr.get("color")
        stroke_hex, sw = None, 0.0
        if stroke and not all(c > 0.99 for c in stroke):
            stroke_hex = "%02X%02X%02X" % (
                int(round(stroke[0] * 255)), int(round(stroke[1] * 255)),
                int(round(stroke[2] * 255)))
            try:
                sw = float(dr.get("width") or 0.5)
            except Exception:
                sw = 0.5
        out.append((tuple(r), fill_hex, stroke_hex, sw))
    return out


def _bg_fill_for(bbox, bg_rects):
    """Innermost (smallest-area) background rect containing the block center.

    Returns ("RRGGBB" fill, "RRGGBB" stroke or None, stroke_width_pt); the
    stroke is inherited from the nearest outer rect when the innermost rect
    itself has none (PDF code boxes: filled inner rect + stroked outer frame)."""
    if not bg_rects:
        return None, None, 0.0
    cx = (bbox[0] + bbox[2]) / 2.0
    cy = (bbox[1] + bbox[3]) / 2.0
    best, best_area = None, None
    for r, fill, stroke, sw in bg_rects:
        if r[0] <= cx <= r[2] and r[1] <= cy <= r[3]:
            area = (r[2] - r[0]) * (r[3] - r[1])
            if best is None or area < best_area:
                best, best_area = (fill, stroke, sw), area
    if best is None:
        return None, None, 0.0
    fill, stroke, sw = best
    if not stroke:
        cands = [((r[2] - r[0]) * (r[3] - r[1]), st, w)
                 for r, _f, st, w in bg_rects
                 if st and r[0] <= cx <= r[2] and r[1] <= cy <= r[3]]
        if cands:
            cands.sort()
            stroke, sw = cands[0][1], cands[0][2]
    return fill, stroke, (sw or 0.5)


def _apply_paragraph_shading(p, color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)


def _apply_paragraph_border(p, color, sides, width_pt=0.5):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    sz = max(2, int(round(width_pt * 8)))
    for side in sides:
        el = OxmlElement("w:%s" % side)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)


def _apply_paragraph_border(p, color, sides, width_pt=0.5):
    """w:pBdr on the given sides (top/bottom/left/right) of one paragraph."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    sz = max(2, int(round(width_pt * 8)))  # eighths of a point
    for side in sides:
        el = OxmlElement("w:%s" % side)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)


def _add_box_gap(doc, info):
    """Shaded empty paragraph filling a blank line inside a code box so the
    background stays continuous (the PDF box fill covers its blank lines)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(info["h_pt"])  # exact line height
    _apply_paragraph_shading(p, info["bg_fill"])
    if info.get("bg_stroke"):
        _apply_paragraph_border(p, info["bg_stroke"], ["left", "right"],
                                info.get("bg_stroke_w") or 0.5)


def convert_digital_page(doc, page, pdf_doc, region, body_size, page_idx, warnings):
    """Convert one digital (text-layer) page. Returns stats dict."""
    stats = {"mode": "digital", "paragraphs": 0, "tables": 0, "images": 0, "chars": 0}
    n, td = page_char_count(page)
    stats["chars"] = n
    if n < 2:
        return stats

    ul_marks = underline_lines(page)
    try:
        tbl = page.find_tables()
        tables = [t for t in tbl.tables if _is_real_table(t)]
    except Exception as e:
        tables = []
        warnings.append("page %d: table detection failed: %s" % (page_idx + 1, e))
    table_rects = [t.bbox for t in tables]
    bg_rects = page_bg_rects(page)

    # build flow items
    items = []
    for b in td["blocks"]:
        if b.get("type") == 1:
            items.append({"kind": "image", "bbox": b["bbox"], "xref": b.get("xref"),
                          "image": b.get("image")})
        elif b.get("type") == 0:
            lines = [l for l in b["lines"]
                     if not any(center_in_rect(((l["bbox"][0] + l["bbox"][2]) / 2,
                                                (l["bbox"][1] + l["bbox"][3]) / 2), r, pad=2.0)
                                for r in table_rects)]
            if not lines:
                continue
            block = dict(b)
            block["lines"] = lines
            info = para_info_from_block(block, region, page, ul_marks, page.rect)
            if info:
                info["text"] = "".join(s["text"] for l in lines for s in l["spans"])
                (info["bg_fill"], info["bg_stroke"],
                 info["bg_stroke_w"]) = _bg_fill_for(info["bbox"], bg_rects)
                items.append({"kind": "para", "info": info, "bbox": info["bbox"]})
    for t in tables:
        items.append({"kind": "table", "table": t, "bbox": t.bbox})

    # sort by y with 2-column awareness
    items.sort(key=lambda it: (it["bbox"][1], it["bbox"][0]))
    for i in range(len(items) - 1):
        a, b = items[i], items[i + 1]
        ov = rect_overlap_v(a["bbox"], b["bbox"])
        minh = min(a["bbox"][3] - a["bbox"][1], b["bbox"][3] - b["bbox"][1])
        if minh > 0 and ov > 0.5 * minh and a["bbox"][0] > b["bbox"][0] + 12:
            items[i], items[i + 1] = b, a

    # collect line-level spans for table cell styling
    page_spans = []
    for b in td["blocks"]:
        if b.get("type") != 0:
            continue
        for l in b["lines"]:
            for s in l["spans"]:
                if s.get("text", "").strip():
                    page_spans.append((l["bbox"], s["text"], {
                        "font": s["font"], "bold": bool(s["flags"] & 16)}))

    # paragraph gaps
    prev_bottom = None
    for it in items:
        if it["kind"] == "para":
            info = it["info"]
            if prev_bottom is not None:
                gap = info["bbox"][1] - prev_bottom
                if gap > 1.3 * info["size"]:
                    info["space_before"] = min(36.0, gap - info["size"] * 0.5)
            prev_bottom = info["bbox"][3]
        else:
            prev_bottom = it["bbox"][3]

    # group consecutive same-fill paragraphs into code boxes: one outer frame
    # per group (top border on the first line, bottom on the last, left+right
    # on every line), so the docx shows the PDF's stroked box, and fill
    # in-box gaps with shaded empty paragraphs so the background stays
    # continuous like the PDF.
    def _close_box(g):
        for j, it in enumerate(g):
            sides = []
            if j == 0:
                sides.append("top")
                it["info"]["box_first"] = True
            if j == len(g) - 1:
                sides.append("bottom")
            sides += ["left", "right"]
            it["info"]["box_sides"] = sides

    group = []
    for it in items:
        f = it["info"].get("bg_fill") if it["kind"] == "para" else None
        if f and group and it["info"]["bg_fill"] == group[-1]["info"]["bg_fill"]:
            group.append(it)
        elif f:
            if group:
                _close_box(group)
            group = [it]
        else:
            if group:
                _close_box(group)
            group = []
    if group:
        _close_box(group)

    box_items = []
    for it in items:
        if it["kind"] == "para" and it["info"].get("box_sides") \
                and not it["info"].get("box_first"):
            sb = it["info"].get("space_before") or 0.0
            if sb >= 8.0:
                box_items.append({"kind": "boxgap", "info": {
                    "bg_fill": it["info"]["bg_fill"],
                    "bg_stroke": it["info"].get("bg_stroke"),
                    "bg_stroke_w": it["info"].get("bg_stroke_w"),
                    "h_pt": sb,
                }})
                it["info"]["space_before"] = 0.0
        box_items.append(it)
    items = box_items

    for it in items:
        if it["kind"] == "para":
            add_digital_paragraph(doc, it["info"], body_size, page)
            stats["paragraphs"] += 1
        elif it["kind"] == "boxgap":
            _add_box_gap(doc, it["info"])
            stats["paragraphs"] += 1
        elif it["kind"] == "table":
            res = add_digital_table(doc, it["table"], page_spans, body_size)
            if res:
                stats["tables"] += 1
        elif it["kind"] == "image":
            if (it["bbox"][2] - it["bbox"][0]) > 24 and (it["bbox"][3] - it["bbox"][1]) > 24:
                if add_digital_image(doc, pdf_doc, it, region):
                    stats["images"] += 1
    return stats


# ---------------------------------------------------------------------------
# Scan (OCR) page conversion
# ---------------------------------------------------------------------------
def convert_scan_page(doc, page, page_img_path, body_size, page_idx, warnings, ocr_engine,
                      body_is_global=False):
    from scanocr import analyze_page, add_scan_page
    stats = {"mode": "scan", "paragraphs": 0, "tables": 0, "images": 0, "chars": 0}
    data = analyze_page(page_img_path, ocr_engine)
    stats["chars"] = sum(len(l["text"]) for l in data["lines"])
    res = add_scan_page(doc, data, page_img_path, page.rect, body_size, page_idx, warnings,
                        body_is_global=body_is_global)
    stats.update(res)
    return stats


# ---------------------------------------------------------------------------
# Body size estimation
# ---------------------------------------------------------------------------
def estimate_body_size(pdf_doc, pages, force_scan):
    sizes = Counter()
    for pi in pages:
        page = pdf_doc[pi]
        n, td = page_char_count(page)
        if force_scan or n < 20:
            continue
        for b in td["blocks"]:
            if b.get("type") != 0:
                continue
            for l in b["lines"]:
                for s in l["spans"]:
                    if s.get("text", "").strip():
                        sizes[round(s["size"], 1)] += len(s["text"])
    if sizes:
        return sizes.most_common(1)[0][0]
    return 10.5


# ---------------------------------------------------------------------------
# Section / page geometry
# ---------------------------------------------------------------------------
def setup_section(doc, pdf_doc, pages):
    p0 = pdf_doc[pages[0]]
    sec = doc.sections[0]
    W = p0.rect.width
    H = p0.rect.height
    sec.page_width = Pt(W)
    sec.page_height = Pt(H)
    # margins from the content extremes (min/max) across text and images so the
    # docx text column matches the PDF's; a per-block median collapses to huge
    # margins whenever many short (code) lines are present.
    xs0, xs1, ys0, ys1 = [], [], [], []
    for pi in pages:
        page = pdf_doc[pi]
        d = page.get_text("dict")
        for b in d["blocks"]:
            if b.get("type") in (0, 1):
                bb = b["bbox"]
                if bb[2] - bb[0] < 2 or bb[3] - bb[1] < 2:
                    continue
                xs0.append(bb[0]); xs1.append(bb[2])
                ys0.append(bb[1]); ys1.append(bb[3])
    if xs0:
        sec.left_margin = Pt(min(max(18.0, min(xs0)), 144.0))
        sec.right_margin = Pt(min(max(18.0, W - max(xs1)), 144.0))
        sec.top_margin = Pt(min(max(18.0, min(ys0)), 144.0))
        sec.bottom_margin = Pt(min(max(18.0, H - max(ys1)), 144.0))


def text_region(pdf_doc, pages):
    xs0, xs1, ys0, ys1 = [], [], [], []
    for pi in pages:
        d = pdf_doc[pi].get_text("dict")
        for b in d["blocks"]:
            if b.get("type") == 0:
                bb = b["bbox"]
                xs0.append(bb[0]); xs1.append(bb[2])
                ys0.append(bb[1]); ys1.append(bb[3])
    if xs0:
        return (min(xs0), min(ys0), max(xs1), max(ys1))
    p0 = pdf_doc[pages[0]]
    return (36, 36, p0.rect.width - 36, p0.rect.height - 36)


# ---------------------------------------------------------------------------
# Word page rendering for verification
# ---------------------------------------------------------------------------
def find_soffice():
    cands = [
        "soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in cands:
        if shutil.which(c):
            return c
        if os.path.exists(c):
            return c
    return None


def render_word_soffice(soffice, docx_path, assets_dir, dpi):
    tmp = tempfile.mkdtemp(prefix="pdf2w_lo_")
    subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", tmp, docx_path],
                   timeout=600, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                   check=True)
    out = os.path.join(tmp, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    d = fitz.open(out)
    count = 0
    for i, pg in enumerate(d):
        pix = pg.get_pixmap(dpi=dpi)
        pix.save(os.path.join(assets_dir, "word_p%d.jpg" % i), jpg_quality=85)
        count += 1
    d.close()
    return count


def render_word_pil(docx_path, assets_dir, dpi):
    from preview import render_docx_pages
    return render_docx_pages(docx_path, assets_dir, dpi)


def _norm(s):
    return re.sub(r"\s+", "", s).lower()


def align_word_pages(pdf_doc, word_page_texts):
    """Map each PDF page to the Word page where its content starts.

    The docx starts every PDF page at an explicit page break, but a page's
    content can overflow into one or two extra Word pages. Content order is
    preserved, so the assignment is monotonic and page 0 is anchored.
    Returns {pdf_idx: word_idx} with the true start page per PDF page."""
    pdf_texts = [_norm(pdf_doc[i].get_text()) for i in range(len(pdf_doc))]

    def grams(s, n=4):
        return set(s[i:i + n] for i in range(max(0, len(s) - n + 1)))

    pdf_grams = [grams(t) for t in pdf_texts]
    word_norm = [_norm(t) for t in word_page_texts]
    word_grams = [grams(t) for t in word_norm]

    def score(pg, wg):
        if not pg:
            return 0.0
        return len(pg & wg) / float(len(pg))

    mapping = {0: 0}
    for pi in range(1, len(pdf_texts)):
        best, best_s = pi, -1.0
        for wi in range(mapping[pi - 1] + 1, len(word_page_texts)):
            s = score(pdf_grams[pi], word_grams[wi])
            if s > best_s:
                best, best_s = wi, s
        mapping[pi] = min(best, len(word_page_texts) - 1)
    return mapping


def compose_word_pages(mapping, assets_dir, n_word_pages, n_pdf_pages):
    """Build word_p<pdf_idx>.jpg by stacking the consecutive Word pages that
    hold PDF page <pdf_idx>'s content (start page .. page before the next
    PDF page starts). Originals are stashed as word_src_<i>.jpg."""
    for wi in range(n_word_pages):
        src = os.path.join(assets_dir, "word_p%d.jpg" % wi)
        if os.path.exists(src):
            os.replace(src, os.path.join(assets_dir, "word_src_%d.jpg" % wi))

    spans = []
    for pi in range(n_pdf_pages):
        start = mapping.get(pi, pi)
        if pi + 1 in mapping:
            end = mapping[pi + 1]
        else:
            end = n_word_pages  # last page: keep until content runs out
        spans.append((max(start, 0), min(end, n_word_pages)))

    from PIL import Image as _Im
    moved = []
    for pi, (s, e) in enumerate(spans):
        if s >= e:
            continue
        imgs = []
        for wi in range(s, e):
            p = os.path.join(assets_dir, "word_src_%d.jpg" % wi)
            if os.path.exists(p):
                try:
                    imgs.append(_Im.open(p).convert("RGB"))
                except Exception:
                    pass
        if not imgs:
            continue
        w = max(im.width for im in imgs)
        h = sum(im.height for im in imgs)
        canvas = _Im.new("RGB", (w, h), "white")
        y = 0
        for im in imgs:
            canvas.paste(im, (0, y))
            y += im.height
        dst = os.path.join(assets_dir, "word_p%d.jpg" % pi)
        canvas.save(dst, quality=85)
        if e - s > 1:
            moved.append((pi, (s, e)))
    return moved


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--pdf", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--mode", default="auto", choices=["auto", "digital", "scan"])
    ap.add_argument("--assets", default=None)
    ap.add_argument("--dpi", type=int, default=140)
    ap.add_argument("--render", default="auto", choices=["auto", "pil", "soffice"])
    args = ap.parse_args()

    warnings = []
    assets_dir = args.assets or os.path.join(os.path.dirname(os.path.abspath(args.out)), "pdf2w_assets")
    os.makedirs(assets_dir, exist_ok=True)

    pdf_doc = fitz.open(args.pdf)
    pages = list(range(len(pdf_doc)))
    if not pages:
        print("PDF2WORD_JSON " + json.dumps({"ok": False, "error": "empty pdf"}))
        return 1

    # per-page mode
    page_modes = {}
    for pi in pages:
        if args.mode == "digital":
            page_modes[pi] = "digital"
        elif args.mode == "scan":
            page_modes[pi] = "scan"
        else:
            n, _ = page_char_count(pdf_doc[pi])
            page_modes[pi] = "scan" if n < 20 else "digital"
    n_scan = sum(1 for v in page_modes.values() if v == "scan")
    n_dig = len(pages) - n_scan
    if n_scan and n_dig:
        warnings.append("mixed document: %d digital page(s), %d scanned page(s)" % (n_dig, n_scan))
    if not n_scan and not n_dig:
        warnings.append("empty pages")

    body_size = estimate_body_size(pdf_doc, pages, args.mode == "scan")
    region = text_region(pdf_doc, pages)

    doc = Document()
    setup_section(doc, pdf_doc, pages)
    # Normal style defaults
    try:
        ns = doc.styles["Normal"]
        ns.font.size = Pt(body_size)
        ascii_f, ea_f = map_font(None)
        set_run_fonts_on_style(ns, ascii_f, ea_f)
    except Exception:
        pass

    ocr_engine = None
    if n_scan:
        try:
            from scanocr import OcrEngine
            ocr_engine = OcrEngine.get()
        except Exception as e:
            warnings.append("OCR engine init failed: %r" % e)

    page_stats = []
    _scan_votes = Counter()
    for idx, pi in enumerate(pages):
        if idx > 0:
            pb = doc.add_paragraph()
            pb.paragraph_format.space_before = Pt(0)
            pb.paragraph_format.space_after = Pt(0)
            r = pb.add_run()
            r.add_break(WD_BREAK.PAGE)
        page = pdf_doc[pi]
        mode = page_modes[pi]
        if mode == "scan":
            # render page to image for OCR + verification
            pix = page.get_pixmap(dpi=args.dpi)
            page_img = os.path.join(assets_dir, "ocr_p%d.png" % pi)
            pix.save(page_img)
            st = convert_scan_page(doc, page, page_img, body_size, pi, warnings, ocr_engine,
                                   body_is_global=bool(_scan_votes))
            try:
                os.unlink(page_img)
            except Exception:
                pass
            # accumulate page votes into a document-level running mode so
            # body size stays consistent across pages
            for k, v in (st.get("body_votes") or {}).items():
                _scan_votes[float(k)] += v
            if _scan_votes:
                body_size = _scan_votes.most_common(1)[0][0]
        else:
            st = convert_digital_page(doc, page, pdf_doc, region, body_size, pi, warnings)
        st["page"] = pi
        page_stats.append(st)

    # reflect the final (document-level) body size in the Normal style
    if n_scan and _scan_votes:
        try:
            doc.styles["Normal"].font.size = Pt(body_size)
        except Exception:
            pass

    out_dir = os.path.dirname(os.path.abspath(args.out))
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(args.out)
    warnings.append("body font size estimated at %.1fpt" % body_size)

    # render PDF pages for verification
    for pi in pages:
        pix = pdf_doc[pi].get_pixmap(dpi=args.dpi)
        pix.save(os.path.join(assets_dir, "pdf_p%d.jpg" % pi), jpg_quality=85)

    # render Word pages
    renderer = "pil"
    wpages = 0
    if args.render in ("auto", "soffice"):
        soffice = find_soffice()
        if soffice:
            try:
                wpages = render_word_soffice(soffice, args.out, assets_dir, args.dpi)
                renderer = "soffice"
            except Exception as e:
                warnings.append("soffice render failed, falling back to PIL preview: %r" % e)
                wpages = 0
    if wpages == 0:
        try:
            wpages, word_page_texts = render_word_pil(args.out, assets_dir, args.dpi)
            renderer = "pil"
            page_map = align_word_pages(pdf_doc, word_page_texts)
            compose_word_pages(page_map, assets_dir, wpages, len(pages))
            warnings.append("word pages aligned to pdf pages: %s" % page_map)
        except Exception as e:
            warnings.append("PIL preview render failed: %r" % e)
            wpages = 0

    result = {
        "ok": True,
        "docx": os.path.abspath(args.out),
        "assets": os.path.abspath(assets_dir),
        "pages": page_stats,
        "n_pages": len(pages),
        "n_scan": n_scan,
        "renderer": renderer,
        "warnings": warnings,
    }
    with open(os.path.join(assets_dir, "convert.json"), "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=1)
    print("PDF2WORD_JSON " + json.dumps(result, ensure_ascii=True))
    return 0


def set_run_fonts_on_style(style, ascii_f, ea_f):
    style.font.name = ascii_f
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), ea_f)


if __name__ == "__main__":
    try:
        sys.exit(main())
    except SystemExit:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        print("PDF2WORD_JSON " + json.dumps({"ok": False, "error": repr(e)}, ensure_ascii=True))
        sys.exit(2)
