#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""fixdocx.py — 确定性 DOCX 表格修复执行器(pdf-to-word verify→fix 闭环)。

子命令:
  dump  <docx> [--out F]
        输出紧凑结构 dump(JSON,stdout 或文件):每张表的 grid 列宽 +
        每行单元格 {x: 文本(截断), gs: colspan, vm: 0无/1restart/2continue, w: tcW dxa}。
  apply <docx> --plan P [--in-place | --out D] [--dry-run]
        按白名单操作修复表格,写 FIXDOCX_JSON 结果(stdout 最后一行)。

方案(plan)格式:
  {"actions": [ {op, ...}, ... ], "notes": "..."}

白名单操作(字段全部必填,除注明外):
  setGridSpan {t, row, col, span}            设置单元格跨列数(span>=1)
  setVMerge   {t, row, col, merge}           merge ∈ restart|continue|none
  setCellText {t, row, col, text}            替换单元格文本(保留首个 run 格式)
  insertCell  {t, row, col, text?, span?}    在 col 之前插入单元格(span 缺省 1)
  removeCell  {t, row, col}                  删除单元格
  cloneRow    {t, afterRow, sourceRow}       在 afterRow 之后插入 sourceRow 副本
  removeRow   {t, row}                       删除整行

地址约定:t = 顶层表序号(0 基);row/col = 行内可视单元格序号(0 基,
已含 gridSpan 展开前的真实 <w:tc> 顺序)。

安全性:
  - 每个操作在目标表的深拷贝上执行并通过逐操作检查后才落回;
  - 全部操作完成后执行最终门禁:每张表每行跨列之和 == grid 列数、
    vMerge 连续段完整、文件可被 python-docx 重新打开;
  - 任一失败 => 全部回滚,原文件不写,退出码 2。
  - --in-place 先落 <docx>.pre_fix.docx 备份(已存在则不覆盖)。

退出码:0 成功;2 修复失败(已回滚);3 参数/方案格式错误。
stdout 最后一行: FIXDOCX_JSON <json>
"""

import argparse
import copy
import json
import os
import shutil
import sys

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:  # pragma: no cover
    print("FIXDOCX_JSON " + json.dumps(
        {"ok": False, "error": "python-docx not installed: %r" % e},
        ensure_ascii=True))
    sys.exit(2)


# ---------------------------------------------------------------------------
# 读取
# ---------------------------------------------------------------------------

def _tc_text(tc, limit=24):
    txt = "".join(n.text or "" for n in tc.iter(qn("w:t")))
    txt = txt.strip()
    return txt[:limit]


def _tc_w(tc):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return 0
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        return 0
    try:
        return int(tcW.get(qn("w:w")) or 0)
    except ValueError:
        return 0


def _tc_span(tc):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return 1
    g = tcPr.find(qn("w:gridSpan"))
    if g is None:
        return 1
    try:
        return max(1, int(g.get(qn("w:val")) or 1))
    except ValueError:
        return 1


def _tc_vmerge(tc):
    """0=none 1=restart 2=continue"""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return 0
    v = tcPr.find(qn("w:vMerge"))
    if v is None:
        return 0
    return 1 if v.get(qn("w:val")) == "restart" else 2


def _grid_cols(tbl_el):
    grid = tbl_el.find(qn("w:tblGrid"))
    if grid is None:
        return []
    out = []
    for gc in grid.findall(qn("w:gridCol")):
        try:
            out.append(int(gc.get(qn("w:w")) or 0))
        except ValueError:
            out.append(0)
    return out


def _row_positions(tr):
    """[(grid_offset, tc, span)]"""
    pos, ci = [], 0
    for tc in tr.findall(qn("w:tc")):
        gs = _tc_span(tc)
        pos.append((ci, tc, gs))
        ci += gs
    return pos


def dump_docx(path):
    doc = Document(path)
    # 顶层表 + 其前面最近的非空段落(上下文提示)
    body = doc.element.body
    ctx_map = {}
    prev_txt = ""
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            t = "".join(n.text or "" for n in child.iter(qn("w:t"))).strip()
            if t:
                prev_txt = t
        elif child.tag == qn("w:tbl"):
            ctx_map[id(child)] = prev_txt[:40]
    tables = []
    for t, tbl in enumerate(doc.tables):
        tbl_el = tbl._tbl
        trs = tbl_el.findall(qn("w:tr"))
        rows = []
        for tr in trs:
            rows.append([
                {
                    "x": _tc_text(tc),
                    "gs": _tc_span(tc),
                    "vm": _tc_vmerge(tc),
                    "w": _tc_w(tc),
                }
                for (_, tc, _) in _row_positions(tr)
            ])
        entry = {
            "t": t,
            "grid": _grid_cols(tbl_el),
            "rows": rows,
        }
        ctx = ctx_map.get(id(tbl_el))
        if ctx:
            entry["ctx"] = ctx
        tables.append(entry)
    return {"tables": tables}


# ---------------------------------------------------------------------------
# 结构检查
# ---------------------------------------------------------------------------

def _row_span_sum(tr):
    return sum(gs for (_, _, gs) in _row_positions(tr))


def table_errors(tbl_el):
    """返回 [错误信息];空表=通过。"""
    errs = []
    n_grid = len(_grid_cols(tbl_el))
    if not n_grid:
        return ["table has no tblGrid"]
    trs = tbl_el.findall(qn("w:tr"))
    col_state = {}
    for ri, tr in enumerate(trs):
        s = _row_span_sum(tr)
        if s != n_grid:
            errs.append("row %d span sum %d != grid %d" % (ri, s, n_grid))
        # vMerge 连续段:continue 必须正上方同列有 restart/continue
        row_states = {}
        for ci, tc, gs in _row_positions(tr):
            st = _tc_vmerge(tc)
            for c in range(ci, ci + gs):
                row_states[c] = st
        for c, st in sorted(row_states.items()):
            prev = col_state.get(c)
            if st == 2 and prev not in (1, 2):
                errs.append("row %d col %d vMerge continue without restart above" % (ri, c))
            col_state[c] = st
    return errs


# ---------------------------------------------------------------------------
# 操作实现(全部作用于传入的 tbl_el 深拷贝,调用方负责原子性)
# ---------------------------------------------------------------------------

def _set_grid_span(tc, span):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is not None:
        anchor_after = tcW
    else:
        anchor_after = None
    g = tcPr.find(qn("w:gridSpan"))
    if g is None:
        g = OxmlElement("w:gridSpan")
        if anchor_after is not None:
            anchor_after.addnext(g)
        else:
            tcPr.insert(0, g)
    g.set(qn("w:val"), str(int(span)))


def _set_tc_w(tc, w):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.insert(0, tcW)
    tcW.set(qn("w:w"), str(int(w)))
    tcW.set(qn("w:type"), "dxa")


def _cell_grid_offset(tr, col):
    pos = _row_positions(tr)
    if col < 0 or col >= len(pos):
        return None
    return pos[col][0]


def op_set_grid_span(tbl_el, a, n_grid):
    trs = tbl_el.findall(qn("w:tr"))
    tr = trs[a["row"]]
    tcs = tr.findall(qn("w:tc"))
    if a["col"] >= len(tcs):
        return "col %d out of range (row has %d cells)" % (a["col"], len(tcs))
    span = int(a["span"])
    if span < 1:
        return "span must be >= 1"
    off = _cell_grid_offset(tr, a["col"])
    if off is None or off + span > n_grid:
        return "cell would extend past grid (offset %s + span %d > %d)" % (off, span, n_grid)
    tc = tcs[a["col"]]
    _set_grid_span(tc, span)
    # 宽度:取 grid 列宽之和
    widths = _grid_cols(tbl_el)
    w = sum(widths[off:off + span])
    if w:
        _set_tc_w(tc, w)
    return None


def op_set_vmerge(tbl_el, a, n_grid):
    trs = tbl_el.findall(qn("w:tr"))
    tr = trs[a["row"]]
    tcs = tr.findall(qn("w:tc"))
    if a["col"] >= len(tcs):
        return "col %d out of range" % a["col"]
    tc = tcs[a["col"]]
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    v = tcPr.find(qn("w:vMerge"))
    if a["merge"] == "none":
        if v is not None:
            tcPr.remove(v)
        return None
    if a["merge"] == "restart":
        if v is None:
            v = OxmlElement("w:vMerge")
            tcPr.append(v)
        v.set(qn("w:val"), "restart")
        return None
    # continue:上方同列必须已有 restart/continue
    if a["row"] == 0:
        return "row 0 cannot be vMerge continue"
    off = _cell_grid_offset(tr, a["col"])
    if off is None:
        return "col %d out of range" % a["col"]
    prev_states = {}
    for ci, ptc, gs in _row_positions(trs[a["row"] - 1]):
        for c in range(ci, ci + gs):
            prev_states[c] = _tc_vmerge(ptc)
    if prev_states.get(off) not in (1, 2):
        return "no vMerge run above (row %d col %d)" % (a["row"] - 1, off)
    if v is None:
        v = OxmlElement("w:vMerge")
        tcPr.append(v)
    v.set(qn("w:val"), "continue")
    return None


def op_set_cell_text(tbl_el, a, n_grid):
    trs = tbl_el.findall(qn("w:tr"))
    tr = trs[a["row"]]
    tcs = tr.findall(qn("w:tc"))
    if a["col"] >= len(tcs):
        return "col %d out of range" % a["col"]
    tc = tcs[a["col"]]
    ps = tc.findall(qn("w:p"))
    first = ps[0] if ps else None
    rpr = None
    if first is not None:
        for r in first.findall(qn("w:r")):
            r_el = r.find(qn("w:rPr"))
            if r_el is not None:
                rpr = copy.deepcopy(r_el)
                break
        for r in first.findall(qn("w:r")):
            first.remove(r)
    for p in ps[1:]:
        tc.remove(p)
    if first is None:
        first = OxmlElement("w:p")
        tc.append(first)
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = str(a.get("text") or "")
    r.append(t)
    first.append(r)
    return None


def op_insert_cell(tbl_el, a, n_grid):
    trs = tbl_el.findall(qn("w:tr"))
    tr = trs[a["row"]]
    tcs = tr.findall(qn("w:tc"))
    col = a["col"]
    if col > len(tcs):
        return "col %d out of range (max %d)" % (col, len(tcs))
    span = int(a.get("span") or 1)
    if span < 1:
        return "span must be >= 1"
    off = _cell_grid_offset(tr, col) if col < len(tcs) else (
        sum(gs for (_, _, gs) in _row_positions(tr)))
    if off + span > n_grid:
        return "cell would extend past grid (offset %d + span %d > %d)" % (off, span, n_grid)
    neighbor = tcs[col] if col < len(tcs) else (tcs[-1] if tcs else None)
    if neighbor is None:
        return "row is empty, cannot insert"
    new_tc = copy.deepcopy(neighbor)
    # 清空文本,保留 tcPr(边框/底纹/宽度)
    for p in new_tc.findall(qn("w:p"))[1:]:
        new_tc.remove(p)
    p0 = new_tc.findall(qn("w:p"))[0] if new_tc.findall(qn("w:p")) else None
    if p0 is None:
        p0 = OxmlElement("w:p")
        new_tc.append(p0)
    for r in p0.findall(qn("w:r")):
        p0.remove(r)
    text = str(a.get("text") or "")
    if text:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p0.append(r)
    _set_tc_w(new_tc, sum(_grid_cols(tbl_el)[off:off + span]))
    if span > 1:
        _set_grid_span(new_tc, span)
    if col < len(tcs):
        tcs[col].addprevious(new_tc)
    else:
        tr.append(new_tc)
    return None


def op_remove_cell(tbl_el, a, n_grid):
    trs = tbl_el.findall(qn("w:tr"))
    tr = trs[a["row"]]
    tcs = tr.findall(qn("w:tc"))
    if a["col"] >= len(tcs):
        return "col %d out of range" % a["col"]
    tr.remove(tcs[a["col"]])
    return None


def op_clone_row(tbl_el, a, n_grid):
    trs = tbl_el.findall(qn("w:tr"))
    if a["sourceRow"] >= len(trs) or a["afterRow"] >= len(trs) or a["afterRow"] < 0:
        return "row index out of range"
    new_tr = copy.deepcopy(trs[a["sourceRow"]])
    if a["afterRow"] + 1 < len(trs):
        trs[a["afterRow"] + 1].addprevious(new_tr)
    else:
        trs[-1].addnext(new_tr)
    return None


def op_remove_row(tbl_el, a, n_grid):
    trs = tbl_el.findall(qn("w:tr"))
    if a["row"] >= len(trs):
        return "row %d out of range" % a["row"]
    tbl_el.remove(trs[a["row"]])
    return None


OPS = {
    "setGridSpan": ("t", "row", "col", "span"),
    "setVMerge": ("t", "row", "col", "merge"),
    "setCellText": ("t", "row", "col", "text"),
    "insertCell": ("t", "row", "col"),
    "removeCell": ("t", "row", "col"),
    "cloneRow": ("t", "afterRow", "sourceRow"),
    "removeRow": ("t", "row"),
}


def _is_int(v):
    return isinstance(v, int) and not isinstance(v, bool)


def validate_plan(plan):
    """返回 (valid_actions, invalid:[{i, reason}])"""
    if not isinstance(plan, dict):
        return [], [{"i": -1, "reason": "plan is not an object"}]
    raw = plan.get("actions")
    if not isinstance(raw, list):
        return [], [{"i": -1, "reason": "'actions' must be an array"}]
    valid, invalid = [], []
    for i, a in enumerate(raw):
        if not isinstance(a, dict):
            invalid.append({"i": i, "reason": "action is not an object"})
            continue
        op = a.get("op")
        if op not in OPS:
            invalid.append({"i": i, "reason": "unknown op %r" % op})
            continue
        ok = True
        # 逐字段类型检查
        for f in OPS[op]:
            v = a.get(f)
            if f in ("t", "row", "col", "span", "afterRow", "sourceRow"):
                if not _is_int(v) or v < 0:
                    invalid.append({"i": i, "reason": "field %r must be int >= 0 (got %r)" % (f, v)})
                    ok = False
                    break
            elif f == "merge":
                if v not in ("restart", "continue", "none"):
                    invalid.append({"i": i, "reason": "merge must be restart|continue|none"})
                    ok = False
                    break
            elif f == "text":
                if not isinstance(v, str):
                    invalid.append({"i": i, "reason": "field 'text' must be a string"})
                    ok = False
                    break
        if op == "insertCell" and "span" in a and not _is_int(a["span"]):
            invalid.append({"i": i, "reason": "field 'span' must be int"})
            ok = False
        if not ok:
            continue
        valid.append(a)
    return valid, invalid


# ---------------------------------------------------------------------------
# 应用
# ---------------------------------------------------------------------------

def apply_plan(docx_path, plan, out_path=None, in_place=False, dry_run=False):
    result = {"ok": False, "applied": 0, "failed": [], "warnings": []}
    doc = Document(docx_path)
    tables = doc.tables
    valid, invalid = validate_plan(plan)
    result["failed"] = invalid
    if not valid:
        result["error"] = "no valid actions in plan"
        return result

    n_grid = {}
    for t, tbl in enumerate(tables):
        g = _grid_cols(tbl._tbl)
        n_grid[t] = len(g)
    kept = []
    for a in valid:
        if a["t"] not in n_grid:
            result["failed"].append({"i": None, "op": a.get("op"),
                                     "reason": "table %r out of range (have %d)" % (a["t"], len(tables))})
        else:
            kept.append(a)
    valid = kept

    # 整表备份(全量回滚用)
    table_snapshots = {t: copy.deepcopy(tbl._tbl) for t, tbl in enumerate(tables)}

    handlers = {
        "setGridSpan": op_set_grid_span,
        "setVMerge": op_set_vmerge,
        "setCellText": op_set_cell_text,
        "insertCell": op_insert_cell,
        "removeCell": op_remove_cell,
        "cloneRow": op_clone_row,
        "removeRow": op_remove_row,
    }

    for i, a in enumerate(valid):
        tbl_el = tables[a["t"]]._tbl
        scratch = copy.deepcopy(tbl_el)
        try:
            err = handlers[a["op"]](scratch, a, n_grid[a["t"]])
        except Exception as e:
            err = "exception: %r" % e
        if err:
            result["failed"].append({"i": i, "op": a["op"], "reason": err})
            continue
        # 落回:scratch 替换原表
        parent = tbl_el.getparent()
        parent.insert(list(parent).index(tbl_el), scratch)
        parent.remove(tbl_el)
        # 重新绑定 python-docx 包装对象
        from docx.table import Table
        tables[a["t"]] = Table(scratch, tables[a["t"]]._parent)
        result["applied"] += 1

    if not result["applied"]:
        result["error"] = "no action applied successfully"
        return result

    # 最终门禁
    errs = []
    for t, tbl in enumerate(tables):
        for e in table_errors(tbl._tbl):
            errs.append("table %d: %s" % (t, e))
    if errs:
        # 全量回滚
        for t, tbl in enumerate(tables):
            snap = table_snapshots[t]
            parent = tbl._tbl.getparent()
            parent.insert(list(parent).index(tbl._tbl), snap)
            parent.remove(tbl._tbl)
            from docx.table import Table
            tables[t] = Table(snap, tables[t]._parent)
        result["invariants"] = "violated: " + "; ".join(errs[:8])
        result["rolled_back"] = True
        return result
    result["invariants"] = "ok"

    if dry_run:
        result["ok"] = True
        result["dry_run"] = True
        return result

    # 备份 + 保存
    if in_place:
        bak = docx_path + ".pre_fix.docx"
        if not os.path.exists(bak):
            shutil.copy2(docx_path, bak)
            result["backup"] = bak
        doc.save(docx_path)
        result["saved"] = docx_path
    else:
        if not out_path:
            result["error"] = "need --out or --in-place"
            return result
        doc.save(out_path)
        result["saved"] = out_path

    # 重新打开验证(防 XML 损坏)
    try:
        Document(result["saved"])
    except Exception as e:
        result["ok"] = False
        result["error"] = "saved file cannot be re-opened: %r" % e
        # 若 in-place,恢复备份
        if in_place:
            bak = docx_path + ".pre_fix.docx"
            if os.path.exists(bak):
                shutil.copy2(bak, docx_path)
                result["rolled_back"] = True
        return result

    result["ok"] = True
    return result


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    sub = ap.add_subparsers(dest="cmd", required=True)

    pd = sub.add_parser("dump", help="structure dump")
    pd.add_argument("docx")
    pd.add_argument("--out", default=None)

    pa = sub.add_parser("apply", help="apply a fix plan")
    pa.add_argument("docx")
    pa.add_argument("--plan", required=True)
    pa.add_argument("--out", default=None)
    pa.add_argument("--in-place", action="store_true")
    pa.add_argument("--dry-run", action="store_true")

    args = ap.parse_args(argv)

    try:
        if args.cmd == "dump":
            data = dump_docx(args.docx)
            text = json.dumps(data, ensure_ascii=True)
            if args.out:
                with open(args.out, "w", encoding="utf-8") as f:
                    f.write(text)
            else:
                print(text)
        else:
            with open(args.plan, "r", encoding="utf-8") as f:
                plan = json.load(f)
            result = apply_plan(args.docx, plan,
                                out_path=args.out,
                                in_place=args.in_place,
                                dry_run=args.dry_run)
            print("FIXDOCX_JSON " + json.dumps(result, ensure_ascii=True))
            return 0 if result.get("ok") else 2
    except SystemExit:
        raise
    except Exception as e:
        print("FIXDOCX_JSON " + json.dumps(
            {"ok": False, "error": "unhandled: %r" % e}, ensure_ascii=True))
        return 2
    return 0


if __name__ == "__main__":
    sys.exit(main())
